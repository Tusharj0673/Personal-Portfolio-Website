import os
import re
import json
import asyncio
import requests
from io import BytesIO
from pathlib import Path
from typing import AsyncGenerator

from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from dotenv import load_dotenv
from groq import Groq
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from pypdf import PdfReader
from docx import Document

# Load .env.local from the root project folder
root_dir = Path(__file__).resolve().parent.parent
load_dotenv(root_dir / ".env.local")

# ==========================================
# 1. SETUP & CONFIGURATION
# ==========================================
limiter = Limiter(key_func=get_remote_address)
app = FastAPI()
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

GROQ_API_KEY = os.getenv("GROK_API_KEY") or os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("Error: API Key missing. Please set GROQ_API_KEY in your .env file.")

client = Groq(api_key=GROQ_API_KEY)
MODEL_NAME = "openai/gpt-oss-120b"


# ==========================================
# 2. GOOGLE DRIVE RESUME LOADER
# ==========================================
def extract_file_id_from_url(url: str) -> str:
    match = re.search(r"/d/([a-zA-Z0-9_-]+)", url)
    if match: return match.group(1)
    match_id = re.search(r"id=([a-zA-Z0-9_-]+)", url)
    if match_id: return match_id.group(1)
    raise ValueError("Invalid Google Drive URL format.")

def load_resume_from_gdrive(shareable_url: str) -> str:
    file_id = extract_file_id_from_url(shareable_url)
    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
    response = requests.get(download_url)
    
    if "download_warning" in response.cookies:
        for key, value in response.cookies.items():
            if key.startswith("download_warning"):
                response = requests.get(download_url, params={"confirm": value})
                break

    if response.status_code != 200:
        raise Exception(f"Failed to download file from Google Drive. Status code: {response.status_code}")

    pdf_file = BytesIO(response.content)
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text() + "\n"
    return text

# ==========================================
# 3. STRUCTURED SCHEMAS & PARSERS
# ==========================================
class Experience(BaseModel):
    company: str | None = Field(default=None)
    role: str | None = Field(default=None)
    duration: str | None = Field(default=None)
    description: str | None = Field(default=None)
    skills_used: list[str] = Field(default_factory=list)

class ProjectDetail(BaseModel):
    title: str = Field(description="Name of the project")
    tech_stack: list[str] = Field(default_factory=list, description="Technologies, frameworks, and databases used in this project")
    description: str | None = Field(default=None, description="Overview of what the project does")


class ResumeSchema(BaseModel):
    name: str | None = None
    skills: list[str] = Field(default_factory=list)
    experiences: list[Experience] = Field(default_factory=list)
    education: list[str] = Field(default_factory=list)
    projects: list[ProjectDetail] = Field(default_factory=list) # <--- Updated to structured 
    certifications: list[str] = Field(default_factory=list)

def parse_resume_to_json(resume_text: str) -> ResumeSchema:
    schema_json = ResumeSchema.model_json_schema()
    system_prompt = f"""
You are an expert resume parsing engine.
You MUST respond strictly in valid JSON matching the following schema:
{json.dumps(schema_json, indent=2)}
Rules: Include internships inside experiences. Do not invent details.
"""
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Parse this resume:\n\n{resume_text}"}
    ]
    response = client.chat.completions.create(
        model=MODEL_NAME, messages=messages, temperature=0.0, response_format={"type": "json_object"}
    )
    return ResumeSchema(**json.loads(response.choices[0].message.content))

GDRIVE_RESUME_URL = os.getenv("RESUME_GDRIVE_URL")
if not GDRIVE_RESUME_URL:
    raise ValueError("Error: RESUME_GDRIVE_URL is missing from your .env.local file.")

print("Fetching and parsing resume...")
raw_resume_text = load_resume_from_gdrive(GDRIVE_RESUME_URL)
structured_resume: ResumeSchema = parse_resume_to_json(raw_resume_text)
print(f"✅ Loaded live resume for: {structured_resume.name}")

# ==========================================
# 4. CUSTOM QA DATASET LOADER
# ==========================================
QA_FILE_PATH = Path(__file__).resolve().parent / "custom_qa.json"

def load_custom_qa() -> str:
    if not QA_FILE_PATH.exists(): return "No custom QA dataset provided."
    try:
        with open(QA_FILE_PATH, "r", encoding="utf-8") as f:
            qa_list = json.load(f)
        formatted_qa = ""
        for idx, item in enumerate(qa_list, 1):
            formatted_qa += f"Q{idx}: {item['question']}\nA{idx}: {item['answer']}\n\n"
        return formatted_qa
    except Exception as e:
        return "No custom QA dataset provided."

CUSTOM_QA_KNOWLEDGE = load_custom_qa()

# ==========================================
# 5. CHAT ENGINE WITH DYNAMIC PERSONA
# ==========================================
class PortfolioChatbot:
    def __init__(self, resume_data: ResumeSchema):
        self.candidate_name = resume_data.name or "Tushar Jain"
        self.resume_json_str = resume_data.model_dump_json(indent=2)
        
        # Highly optimized system prompt for Persona and Formatting
        self.system_prompt = f"""
You are the dedicated AI Portfolio Assistant representing {self.candidate_name}.
Your primary role is to advocate for {self.candidate_name} to recruiters by being conversational, interactive, and structured, while remaining strictly an AI assistant speaking in the third person.

--- CANDIDATE RESUME DATA (JSON) ---
{self.resume_json_str}

--- PRE-TRAINED KNOWLEDGE & PERSPECTIVES ---
{CUSTOM_QA_KNOWLEDGE}

--- CORE PERSONA & COMMUNICATION RULES ---

1. MANDATORY OPENING ANCHOR:
   - When asked about skills: The first line of your response must be:
     "{self.candidate_name}'s skills include:"
   - When asked about projects: The first line of your response must be:
     "{self.candidate_name} has worked on the following projects:"

2. THIRD-PERSON PERSPECTIVE:
   - Always speak about {self.candidate_name} in the third person (e.g., "{self.candidate_name} built", "He is skilled in", "His background").
   - Never use first-person pronouns ("I", "my", "me"). If asked "Why should we hire you?", respond: "You should hire {self.candidate_name} because he..."

3. SKILLS FORMATTING & ISOLATION:
   - Extract skills strictly from the resume's dedicated "skills" section.
   - Do not pull extra libraries or tools from project descriptions into the core skills list.
   - Organize skills into technical categories (e.g., Programming Languages, Web Development, Databases, Frameworks, Containerization & Tools) using Markdown bullet points (`*`) and bold category headers.

4. TECHNICAL COMPLEXITY & CHALLENGE EVALUATION (MANDATORY DIRECTIVE):
   - When asked evaluative or subjective questions (such as "Which project was the most challenging?", "What was his hardest project?", or "What is his most complex work?"), YOU MUST PERFORM A TECHNICAL COMPLEXITY EVALUATION. Never state that this detail is missing or unlisted.
   - Evaluation Criteria: Compare the projects in the provided resume data based on:
     a) Integration of advanced models or specialized APIs.
     b) Multi-tier architecture (e.g., modern frontend + dedicated backend + database + external services).
     c) Data processing pipelines and explainability/performance requirements.
   - Identify the project that best fits these criteria, name it, and present a structured breakdown explaining the specific technical challenges (e.g., multi-technology integration, backend model orchestration, data pipeline management) that made it demanding.

5. CUSTOM QA ALIGNMENT:
   - For questions related to salary/stipend expectations, location preferences, work culture, or personal background covered in the Pre-Trained Knowledge, faithfully reflect the provided points while translating them into third-person assistant speech.

6. PRIVACY GUARDRAIL:
   - If asked for direct personal contact details (phone number or personal email), reply strictly:
     "I am not allowed to disclose direct contact details. Please refer to his downloaded resume for that information."

--- JOB DESCRIPTION ANALYSIS TRIGGER ---
If the user provides a Job Description (either pasted or uploaded), evaluate {self.candidate_name}'s alignment against the requirements and format your response EXACTLY as this markdown list:
- **Candidate Name:** [Name]
- **Matching Skills:** [Comma-separated list]
- **Missing Important Skills:** [Comma-separated list]
- **Experience Requirement Met:** [Yes/No/Partial - Brief reason]
- **Overall Match Percentage:** [0-100%]
- **Final Verdict:** [1-2 sentences on whether they should interview him and why]
"""
        self.history = [{"role": "system", "content": self.system_prompt}]

sessions = {}

class ChatPayload(BaseModel):
    session_id: str
    message: str

async def generate_groq_stream(session_id: str, user_message: str) -> AsyncGenerator[str, None]:
    if session_id not in sessions:
        sessions[session_id] = PortfolioChatbot(structured_resume).history
    
    sessions[session_id].append({"role": "user", "content": user_message})

    try:
        response_stream = client.chat.completions.create(
            model=MODEL_NAME, messages=sessions[session_id], temperature=0.49, stream=True
        )

        full_reply = ""
        for chunk in response_stream:
            content = chunk.choices[0].delta.content
            if content:
                full_reply += content
                yield f"data: {json.dumps({'content': content})}\n\n"
                await asyncio.sleep(0.015)

        sessions[session_id].append({"role": "assistant", "content": full_reply})
        if len(sessions[session_id]) > 11:
            sessions[session_id] = [sessions[session_id][0]] + sessions[session_id][-10:]
        yield "data: [DONE]\n\n"

    except Exception as e:
        yield f"data: {json.dumps({'error': str(e)})}\n\n"

# ==========================================
# 6. ENDPOINTS
# ==========================================
@app.post("/api/chat/stream")
@limiter.limit("15/minute")
async def chat_stream_endpoint(request: Request, payload: ChatPayload):
    return StreamingResponse(
        generate_groq_stream(payload.session_id, payload.message),
        media_type="text/event-stream"
    )

@app.post("/api/chat/upload-jd")
@limiter.limit("5/minute")
async def upload_job_description(request: Request, file: UploadFile = File(...), session_id: str = Form(...)):
    """Extracts text from uploaded PDF/DOCX and injects it into the chat session."""
    try:
        content_bytes = await file.read()
        ext = Path(file.filename).suffix.lower()
        extracted_text = ""

        if ext == ".pdf":
            reader = PdfReader(BytesIO(content_bytes))
            for page in reader.pages:
                if page.extract_text():
                    extracted_text += page.extract_text() + "\n"
        
        elif ext == ".docx":
            document = Document(BytesIO(content_bytes))
            for para in document.paragraphs:
                extracted_text += para.text + "\n"
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            extracted_text += cell.text + "\n"
        else:
            return {"error": "Unsupported file format. Please upload PDF or DOCX."}

        # Ensure session exists
        if session_id not in sessions:
            sessions[session_id] = PortfolioChatbot(structured_resume).history

        # Inject the parsed document into the conversation history
        injection_message = f"I am providing a Job Description document for you to analyze. Please analyze my fit for this role based on your system rules.\n\nDOCUMENT CONTENT:\n{extracted_text}"
        sessions[session_id].append({"role": "user", "content": injection_message})

        return {"message": "Job description successfully uploaded and added to context.", "trigger_analysis": True}

    except Exception as e:
        return {"error": f"Failed to process file: {str(e)}"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)